import io
from django.db.models import Sum
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_ORIENTATION
from api.models import Branch


def create_docx_with_tables(tallons, from_date, to_date):
    # Create a new Word document
    doc = Document()

    # Set the orientation to landscape
    section = doc.sections[0]
    section.orientation = WD_ORIENTATION.LANDSCAPE
    new_width, new_height = section.page_height, section.page_width
    section.page_width = new_width
    section.page_height = new_height

    # Title
    title = doc.add_paragraph(
        f'{from_date.strftime("%Y-%m-%d")} дан {to_date.strftime("%Y-%m-%d")} гача ҳолатига жамият корхона ва '
        'ташкилотларда меҳнат муҳофазаси бўйича олинган огоҳлантириш талонлари тўғрисида\n\n'
        'РЖУ лар бўйича олинган талонлар'
    )
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title.runs[0].bold = True

    # First Table
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run("жадвал 1").italic = True

    # Create the first table
    table1 = doc.add_table(rows=2, cols=18)
    table1.style = 'Table Grid'

    # Define the header data
    branches = Branch.objects.all()
    header_data = [(branch.name, 3) for branch in branches]

    # Add headers
    hdr_cells = table1.rows[0].cells
    col_index = 0
    for header, span in header_data:
        cell = hdr_cells[col_index]
        cell.text = header
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i in range(1, span):
            cell.merge(hdr_cells[col_index + i])
        col_index += span

    # Define sub-headers
    sub_headers = ["№1", "№2", "№3"]
    subhdr_cells = table1.rows[1].cells
    col_index = 0
    for _ in range(len(header_data)):
        for sub_header in sub_headers:
            subhdr_cells[col_index].text = sub_header
            subhdr_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            col_index += 1

    # Calculate data row for each branch
    data_row = []
    for branch in branches:
        branch_tallons = tallons.filter(organization__branch=branch)
        count_1 = branch_tallons.filter(tallon_number=1).count()
        count_2 = branch_tallons.filter(tallon_number=2).count()
        count_3 = branch_tallons.filter(tallon_number=3).count()
        data_row.extend([count_1, count_2, count_3])

    row_cells = table1.add_row().cells
    for col_index, item in enumerate(data_row):
        row_cells[col_index].text = str(item)
        row_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the final row of the first table
    final_row = [sum(data_row[i::3]) for i in range(3)]
    final_row.extend([0] * (len(header_data) * 3 - len(final_row)))
    final_row_cells = table1.add_row().cells
    for col_index, item in enumerate(final_row):
        final_row_cells[col_index].text = str(item)
        final_row_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        for i in range(1, 3):
            if col_index % 3 == 0:
                final_row_cells[col_index + i].merge(final_row_cells[col_index])

    # Add a blank paragraph between tables
    doc.add_paragraph()

    # Second Table Title
    title2 = doc.add_paragraph(f'РЖУ лар бўйича {from_date.strftime("%Y-%m-%d")} дан {to_date.strftime("%Y-%m-%d")} '
                               'учун олинган талонларнинг ушлаб қолинган суммалари')
    title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    title2.runs[0].bold = True

    # Calculate the second table data
    table2_data = [
        [branch.name for branch in branches],
        [tallons.filter(organization__branch=branch).aggregate(Sum('consequence_amount'))['consequence_amount__sum']
         for branch in branches]
    ]

    table2_data[1] = [amount if amount is not None else 0 for amount in table2_data[1]]

    table2_data.append([sum(table2_data[1])])

    # Second Table Title
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run("жадвал 2").italic = True

    # Add the second table
    table2 = doc.add_table(rows=2, cols=len(branches))
    table2.style = 'Table Grid'
    hdr_cells = table2.rows[0].cells
    for i, hdr in enumerate(table2_data[0]):
        hdr_cells[i].text = hdr
        hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the data row
    data_cells = table2.rows[1].cells
    for i, data in enumerate(table2_data[1]):
        data_cells[i].text = str(data) if data else '0'
        data_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add the total row spanning all columns
    row = table2.add_row()
    cell = row.cells[0]
    cell.text = 'ЖАМИ'
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    cell.merge(row.cells[len(branches) - 2])
    row.cells[len(branches) - 1].text = str(table2_data[2][0])
    row.cells[len(branches) - 1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Add a blank paragraph between tables
    doc.add_paragraph()
    doc.add_page_break()
    # Third Table Title
    p = doc.add_paragraph()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
    p.add_run("жадвал 3").italic = True

    # Iterate through branches and organizations
    for branch in branches:
        if branches[0] != branch:
            doc.add_page_break()
        branch_title = doc.add_paragraph(branch.name)
        branch_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        branch_title.runs[0].bold = True

        organizations = branch.branch_organizations.all()
        branch_total_consequence_amount = 0
        if organizations.exists():
            for organization in organizations:
                organization_title = doc.add_paragraph(organization.name)
                organization_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                organization_title.runs[0].bold = True

                # Add a new table for this organization
                table3 = doc.add_table(rows=1, cols=9)
                table3.style = 'Table Grid'

                # Define table headers
                hdr_cells = table3.rows[0].cells
                hdr_cells[0].text = 'Т/р'
                hdr_cells[1].text = 'Ф.И.Ш.'
                hdr_cells[2].text = 'Лавозими'
                hdr_cells[3].text = 'Олинган санаси'
                hdr_cells[4].text = 'Талон рақами'
                hdr_cells[5].text = 'Олинган сабаби'
                hdr_cells[6].text = 'Интизомий жазо буйруқ рақами, санаси ва тури'
                hdr_cells[7].text = 'Интизомий жазо оқибатида олиб қолинган пул миқдори'
                hdr_cells[8].text = 'Изоҳ'

                organization_tallons = tallons.filter(organization=organization)
                organization_total_consequence_amount = 0
                # Populate table with data
                for idx, tallon in enumerate(organization_tallons, 1):
                    row_cells = table3.add_row().cells
                    row_cells[0].text = str(idx)
                    row_cells[1].text = tallon.fullname
                    row_cells[2].text = tallon.position
                    row_cells[3].text = tallon.date_received.strftime("%d.%m.%Y")
                    row_cells[4].text = tallon.tallon_number
                    row_cells[5].text = tallon.reason_received
                    row_cells[6].text = f"Рақам: {tallon.discipline_order}, Сана: {tallon.discipline_order_date.strftime('%d.%m.%Y')}, Тури: {tallon.discipline_type}"
                    consequence_amount = tallon.consequence_amount if tallon.consequence_amount else 0
                    organization_total_consequence_amount += consequence_amount
                    row_cells[7].text = f"{tallon.consequence_amount:.2f}" if tallon.consequence_amount else ""
                    row_cells[8].text = tallon.note or ""
                doc.add_paragraph(f'Жами: {organization_total_consequence_amount:.2f}')
                branch_total_consequence_amount += organization_total_consequence_amount
        else:
            # Add a new table for this branch
            table3 = doc.add_table(rows=1, cols=9)
            table3.style = 'Table Grid'

            # Define table headers
            hdr_cells = table3.rows[0].cells
            hdr_cells[0].text = 'Т/р'
            hdr_cells[1].text = 'Ф.И.Ш.'
            hdr_cells[2].text = 'Лавозими'
            hdr_cells[3].text = 'Олинган санаси'
            hdr_cells[4].text = 'Талон рақами'
            hdr_cells[5].text = 'Олинган сабаби'
            hdr_cells[6].text = 'Интизомий жазо буйруқ рақами, санаси ва тури'
            hdr_cells[7].text = 'Интизомий жазо оқибатида олиб қолинган пул миқдори'
            hdr_cells[8].text = 'Изоҳ'

            branch_tallons = tallons.filter(branch=branch, organization__isnull=True)
            # Populate table with data
            for idx, tallon in enumerate(branch_tallons, 1):
                row_cells = table3.add_row().cells
                row_cells[0].text = str(idx)
                row_cells[1].text = tallon.fullname
                row_cells[2].text = tallon.position
                row_cells[3].text = tallon.date_received.strftime("%d.%m.%Y")
                row_cells[4].text = tallon.tallon_number
                row_cells[5].text = tallon.reason_received
                row_cells[6].text = f"Рақам: {tallon.discipline_order}, Сана: {tallon.discipline_order_date.strftime('%d.%m.%Y')}, Тури: {tallon.discipline_type}"
                consequence_amount = tallon.consequence_amount if tallon.consequence_amount else 0
                branch_total_consequence_amount += consequence_amount
                row_cells[7].text = f"{tallon.consequence_amount:.2f}" if tallon.consequence_amount else ""
                row_cells[8].text = tallon.note or ""
            doc.add_paragraph(f'Жами: {branch_total_consequence_amount:.2f}')
        all_branch_amount = doc.add_paragraph(f'{branch.name} - Жами: {branch_total_consequence_amount:.2f}')
        all_branch_amount.runs[0].bold = True
        all_branch_amount.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    total_consequence_amount = tallons.aggregate(Sum('consequence_amount'))['consequence_amount__sum'] or 0
    all_amount = doc.add_paragraph(f'Умумий жами:  {total_consequence_amount}')
    all_amount.runs[0].bold = True
    all_amount.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

    # Save the document to a BytesIO buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return buffer


# def create_docx_with_tables(tallons, from_date, to_date):
#     # Create a new Word document
#     doc = Document()
#
#     # Set the orientation to landscape
#     section = doc.sections[0]
#     section.orientation = WD_ORIENTATION.LANDSCAPE
#     new_width, new_height = section.page_height, section.page_width
#     section.page_width = new_width
#     section.page_height = new_height
#
#     title = doc.add_paragraph(
#         f'{from_date.strftime("%Y-%m-%d")} дан {to_date.strftime("%Y-%m-%d")} гача ҳолатига жамият корхона ва '
#         'ташкилотларда меҳнат муҳофазаси бўйича олинган огоҳлантириш талонлари тўғрисида\n\n'
#         'РЖУ лар бўйича олинган талонлар'
#     )
#     run = title.runs[0]
#     run.bold = True
#     title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     title.runs[0].font.color.rgb = RGBColor(0, 0, 0)
#
#     # Add the first table title
#     p = doc.add_paragraph()
#     p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
#     run = p.add_run("жадвал 1")
#     run.font.size = Pt(12)
#     run.italic = True
#
#     # Create the first table
#     table1 = doc.add_table(rows=2, cols=18)
#     table1.style = 'Table Grid'
#
#     # Define the header data
#     branches = Branch.objects.all()
#     header_data = [(branch.name, 3) for branch in branches]
#
#     # Add headers
#     hdr_cells = table1.rows[0].cells
#     col_index = 0
#     for header, span in header_data:
#         cell = hdr_cells[col_index]
#         cell.text = header
#         cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         for i in range(1, span):
#             cell.merge(hdr_cells[col_index + i])
#         col_index += span
#
#     # Define sub-headers
#     sub_headers = ["№1", "№2", "№3"]
#     subhdr_cells = table1.rows[1].cells
#     col_index = 0
#     for _ in range(len(header_data)):
#         for sub_header in sub_headers:
#             subhdr_cells[col_index].text = sub_header
#             subhdr_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#             col_index += 1
#
#         # Calculate data row for each branch
#     data_row = []
#     for branch in branches:
#         branch_tallons = tallons.filter(organization__branch=branch)
#         count_1 = branch_tallons.filter(tallon_number=1).count()
#         count_2 = branch_tallons.filter(tallon_number=2).count()
#         count_3 = branch_tallons.filter(tallon_number=3).count()
#         data_row.extend([count_1, count_2, count_3])
#
#     row_cells = table1.add_row().cells
#     for col_index, item in enumerate(data_row):
#         row_cells[col_index].text = str(item)
#         row_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#     # Add the final row of the first table
#     final_row = [sum(data_row[i::3]) for i in range(3)]
#     final_row.extend([0] * (len(header_data) * 3 - len(final_row)))
#     final_row_cells = table1.add_row().cells
#     for col_index, item in enumerate(final_row):
#         final_row_cells[col_index].text = str(item)
#         final_row_cells[col_index].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         for i in range(1, 3):
#             if col_index % 3 == 0:
#                 final_row_cells[col_index + i].merge(final_row_cells[col_index])
#
#     # Add a blank paragraph between tables
#     doc.add_paragraph()
#     title2 = doc.add_paragraph(f'РЖУ лар бўйича {from_date.strftime("%Y-%m-%d")} дан {to_date.strftime("%Y-%m-%d")} '
#                                'учун олинган талонларнинг ушлаб қолинган суммалари')
#     run = title2.runs[0]
#     run.bold = True
#     title2.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     title2.runs[0].font.color.rgb = RGBColor(0, 0, 0)
#
#     # Calculate the second table data
#     table2_data = [
#         [branch.name for branch in branches],
#         [tallons.filter(organization__branch=branch).aggregate(Sum('consequence_amount'))['consequence_amount__sum']
#          for branch in branches]
#     ]
#
#     table2_data[1] = [amount if amount is not None else 0 for amount in table2_data[1]]
#
#     table2_data.append([sum(table2_data[1])])
#
#     # Add the second table title
#     p = doc.add_paragraph()
#     p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
#     run = p.add_run("жадвал 2")
#     run.font.size = Pt(12)
#     run.italic = True
#
#     # Add the second table
#     table2 = doc.add_table(rows=2, cols=len(branches))
#     table2.style = 'Table Grid'
#     hdr_cells = table2.rows[0].cells
#     for i, hdr in enumerate(table2_data[0]):
#         hdr_cells[i].text = hdr
#         hdr_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#     # Add the data row
#     data_cells = table2.rows[1].cells
#     for i, data in enumerate(table2_data[1]):
#         data_cells[i].text = str(data) if data else '0'
#         data_cells[i].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#     # Add the total row spanning all columns
#     row = table2.add_row()
#     cell = row.cells[0]
#     cell.text = 'ЖАМИ'
#     cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#     cell.merge(row.cells[len(branches) - 2])
#     row.cells[len(branches) - 1].text = str(table2_data[2][0])
#     row.cells[len(branches) - 1].paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#
#     # Add a blank paragraph between tables
#     doc.add_paragraph()
#     doc.add_page_break()
#     # Add the third table title
#     p = doc.add_paragraph()
#     p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
#     run = p.add_run("жадвал 3")
#     run.font.size = Pt(12)
#     run.italic = True
#
#     branches = Branch.objects.all()
#     for branch in branches:
#         if branches[0] != branch:
#             doc.add_page_break()
#         branch_title = doc.add_paragraph(branch.name)
#         run = branch_title.runs[0]
#         run.bold = True
#         branch_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#         branch_title.runs[0].font.color.rgb = RGBColor(0, 0, 0)
#
#         organizations = branch.branch_organizations.all()
#         branch_total_consequence_amount = 0
#         if organizations.exists():
#             for organization in organizations:
#                 organization_title = doc.add_paragraph(organization.name)
#                 run = organization_title.runs[0]
#                 run.bold = True
#                 organization_title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
#                 organization_title.runs[0].font.color.rgb = RGBColor(0, 0, 0)
#
#                 # Add a new table for this organization
#                 table3 = doc.add_table(rows=1, cols=9)
#                 table3.style = 'Table Grid'
#
#                 # Define table headers
#                 hdr_cells = table3.rows[0].cells
#                 hdr_cells[0].text = 'Т/р'
#                 hdr_cells[1].text = 'Ф.И.Ш.'
#                 hdr_cells[2].text = 'Лавозими'
#                 hdr_cells[3].text = 'Олинган санаси'
#                 hdr_cells[4].text = 'Талон рақами'
#                 hdr_cells[5].text = 'Олинган сабаби'
#                 hdr_cells[6].text = 'Интизомий жазо буйруқ рақами, санаси ва тури'
#                 hdr_cells[7].text = 'Интизомий жазо оқибатида олиб қолинган пул миқдори'
#                 hdr_cells[8].text = 'Изоҳ'
#
#                 organization_tallons = tallons.filter(organization=organization)
#                 organization_total_consequence_amount = 0
#                 # Populate table with data
#                 for idx, tallon in enumerate(organization_tallons, 1):
#                     row_cells = table3.add_row().cells
#                     row_cells[0].text = str(idx)
#                     row_cells[1].text = tallon.fullname
#                     row_cells[2].text = tallon.position
#                     row_cells[3].text = tallon.date_received.strftime("%d.%m.%Y")
#                     row_cells[4].text = tallon.tallon_number
#                     row_cells[5].text = tallon.reason_received
#                     row_cells[6].text = f"Рақам: {tallon.discipline_order}, Сана: {tallon.discipline_order_date.strftime('%d.%m.%Y')}, Тури: {tallon.discipline_type}"
#                     consequence_amount = tallon.consequence_amount if tallon.consequence_amount else 0
#                     organization_total_consequence_amount += consequence_amount
#                     row_cells[7].text = f"{tallon.consequence_amount:.2f}" if tallon.consequence_amount else ""
#                     row_cells[8].text = tallon.note or ""
#                 doc.add_paragraph(f'Жами: {organization_total_consequence_amount:.2f}')
#                 branch_total_consequence_amount += organization_total_consequence_amount
#         else:
#             # Add a new table for this branch
#             table3 = doc.add_table(rows=1, cols=9)
#             table3.style = 'Table Grid'
#
#             # Define table headers
#             hdr_cells = table3.rows[0].cells
#             hdr_cells[0].text = 'Т/р'
#             hdr_cells[1].text = 'Ф.И.Ш.'
#             hdr_cells[2].text = 'Лавозими'
#             hdr_cells[3].text = 'Олинган санаси'
#             hdr_cells[4].text = 'Талон рақами'
#             hdr_cells[5].text = 'Олинган сабаби'
#             hdr_cells[6].text = 'Интизомий жазо буйруқ рақами, санаси ва тури'
#             hdr_cells[7].text = 'Интизомий жазо оқибатида олиб қолинган пул миқдори'
#             hdr_cells[8].text = 'Изоҳ'
#
#             branch_tallons = tallons.filter(branch=branch, organization__isnull=True)
#             # Populate table with data
#             for idx, tallon in enumerate(branch_tallons, 1):
#                 row_cells = table3.add_row().cells
#                 row_cells[0].text = str(idx)
#                 row_cells[1].text = tallon.fullname
#                 row_cells[2].text = tallon.position
#                 row_cells[3].text = tallon.date_received.strftime("%d.%m.%Y")
#                 row_cells[4].text = tallon.tallon_number
#                 row_cells[5].text = tallon.reason_received
#                 row_cells[6].text = f"Рақам: {tallon.discipline_order}, Сана: {tallon.discipline_order_date.strftime('%d.%m.%Y')}, Тури: {tallon.discipline_type}"
#                 consequence_amount = tallon.consequence_amount if tallon.consequence_amount else 0
#                 branch_total_consequence_amount += consequence_amount
#                 row_cells[7].text = f"{tallon.consequence_amount:.2f}" if tallon.consequence_amount else ""
#                 row_cells[8].text = tallon.note or ""
#             doc.add_paragraph(f'Жами: {branch_total_consequence_amount:.2f}')
#         all_branch_amount = doc.add_paragraph(f'{branch.name} - Жами: {branch_total_consequence_amount:.2f}')
#         run = all_branch_amount.runs[0]
#         run.bold = True
#         all_branch_amount.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#         all_branch_amount.runs[0].font.color.rgb = RGBColor(0, 0, 0)
#
#     all_amount = doc.add_paragraph(f'Умумий жами:  ')
#     run = all_amount.runs[0]
#     run.bold = True
#     all_amount.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
#     all_amount.runs[0].font.color.rgb = RGBColor(0, 0, 0)
#
#     # Save the document to a BytesIO buffer
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#
#     return buffer
#
